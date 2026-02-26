from docx import Document
from datetime import datetime
from pathlib import Path
import json
import subprocess
import os

# Import do módulo web (arquivo separado)
from web import acessar_web

# === 1. Função para gerar a data de hoje formatada ===
def data_hoje_formatada():
    meses = {
        1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril",
        5: "maio", 6: "junho", 7: "julho", 8: "agosto",
        9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro"
    }
    hoje = datetime.now()
    return f"{hoje.day} de {meses[hoje.month]} de {hoje.year}"


def gera_competencia(meses_passados=0):
    hoje = datetime.now()
    competencia = f"{hoje.month - meses_passados}/{hoje.year}"
    if len(competencia) == 6:
        return f"0{competencia}"
    else:
        return competencia


def classificar_mes(data_str):
    try:
        dt = datetime.strptime(data_str.strip(), "%d/%m/%Y").date()
    except ValueError:
        raise ValueError("Data inválida. Use o formato dd/mm/aaaa e uma data real (ex.: 27/01/2026).")

    hoje = datetime.today()
    if (dt.year, dt.month) == (hoje.year, hoje.month):
        return "mes atual"
    else:
        return "mes passado"


# === Define o diretório base (onde está o main.py) ===
BASE_DIR = Path(__file__).resolve().parent
config_path = BASE_DIR / "config.json"

hoje = data_hoje_formatada()
competencia = gera_competencia()

# === Coleta de inputs em um dicionário único ===
dados = {
    "fis": input("Digite o número da folha de deferimento: ").strip(),
    "matricula": input("Digite a matrícula: ").strip(),
    "nome": input("Digite o nome do requerente: ").strip(),
    "cargo": input("Digite o cargo: ").strip(),
    "numero_processo": input("Digite o número do processo: ").strip(),
    "data_deferimento": input("Digite a data de deferimento: ").strip(),
    "inciso": input("Digite o inciso: ").strip(),
    "competencia": competencia,
    "hoje": hoje,
}

# === Monta substituições a partir do dicionário ===
substituicoes = {
    "{{FIS}}": dados["fis"],
    "{{NOME}}": dados["nome"],
    "{{MATRICULA}}": dados["matricula"],
    "{{CARGO}}": dados["cargo"],
    "{{NUMERO_PROCESSO}}": dados["numero_processo"],
    "{{DATA_DEFERIMENTO}}": dados["data_deferimento"],
    "{{COMP}}": dados["competencia"],
    "{{HOJE}}": dados["hoje"],
    "{{INCISO}}": dados["inciso"],
}

# === Escolhe modelo de portaria / grupo ===
if dados["inciso"] in ("I", "II"):
    modelo_portaria = BASE_DIR / "modelos" / "entrada" / "modelo-folha-portaria.docx"
else:
    modelo_portaria = BASE_DIR / "modelos" / "entrada" / "modelo-folha-portaria-com-grupo.docx"
    dados["grupo"] = input("Digite o grupo: ").strip()
    substituicoes["{{GRUPO}}"] = dados["grupo"]

# === Escolhe modelo de informações / retroativo ===
mes_do_deferimento = classificar_mes(dados["data_deferimento"])

if mes_do_deferimento == "mes atual":
    modelo_informacoes = BASE_DIR / "modelos" / "entrada" / "modelo-folha-info2.docx"
else:
    modelo_informacoes = BASE_DIR / "modelos" / "entrada" / "modelo-folha-info-retroativo.docx"

    def calcular_retroativo(salario_base, dia_do_deferimento, inciso):
        dia_do_deferimento = int(dia_do_deferimento[0:2])
        DIAS_POR_MES = 30

        def selecionar_valor_verba(inciso):
            match inciso:
                case "I":
                    return 0.020
                case "II":
                    return 0.025
                case "III":
                    return 0.030
                case "IV":
                    return 0.040
                case "V":
                    return 0.050
                case _:
                    raise ValueError("Verba desconhecida")

        valor_do_adicional = selecionar_valor_verba(inciso)

        dias_para_retroagir = DIAS_POR_MES - dia_do_deferimento + 1

        valor_do_retroativo = round(salario_base / DIAS_POR_MES * valor_do_adicional * dias_para_retroagir, 2)
        return valor_do_retroativo

    salario_base = float(input("Digite o salário base do servidor: ").strip().replace(".", "").replace(",", "."))

    dados["valor_retroativo"] = calcular_retroativo(salario_base, dados["data_deferimento"], dados["inciso"])

    # Correção: removido espaço do placeholder
    substituicoes["{{COMPETENCIA_ANTERIOR}}"] = gera_competencia(1)
    substituicoes["{{VALOR}}"] = dados["valor_retroativo"]

# === Lê config.json ===
if not config_path.exists():
    raise FileNotFoundError("Arquivo config.json não encontrado! Crie um antes de executar.")

with open(config_path, "r", encoding="utf-8") as f:
    config = json.load(f)

saida_informacoes_dir = Path(config["saida_informacoes"])
saida_portaria_dir = Path(config["saida_portaria"])

saida_informacoes_dir.mkdir(parents=True, exist_ok=True)
saida_portaria_dir.mkdir(parents=True, exist_ok=True)

saida_informacoes = saida_informacoes_dir / f"Adicional Conclusão de Curso - {dados['nome']}.docx"
saida_portaria = saida_portaria_dir / f"Adicional por Conclusão de Curso - {dados['nome']}.docx"


# === Funções do DOCX ===
def substituir_texto(paragrafo, mapa):
    if not paragrafo.runs:
        return

    texto_original = "".join(run.text for run in paragrafo.runs)
    texto_novo = texto_original

    for chave, valor in mapa.items():
        texto_novo = texto_novo.replace(chave, str(valor))

    if texto_novo != texto_original:
        paragrafo.runs[0].text = texto_novo
        for run in paragrafo.runs[1:]:
            run.text = ""


def aplicar_substituicoes(caminho_modelo, caminho_saida, mapa):
    doc = Document(caminho_modelo)

    for paragrafo in doc.paragraphs:
        substituir_texto(paragrafo, mapa)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    substituir_texto(paragrafo, mapa)

    doc.save(caminho_saida)

def abrir_no_libreoffice(caminho_arquivo):
    possiveis_caminhos = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
    ]

    soffice = None
    for caminho in possiveis_caminhos:
        if os.path.exists(caminho):
            soffice = caminho
            break

    if soffice is None:
        raise FileNotFoundError("❌ LibreOffice não encontrado no sistema.")

    subprocess.Popen([soffice, caminho_arquivo])

'''
aplicar_substituicoes(modelo_informacoes, saida_informacoes, substituicoes)
aplicar_substituicoes(modelo_portaria, saida_portaria, substituicoes)
abrir_no_libreoffice(str(saida_informacoes))
abrir_no_libreoffice(str(saida_portaria))
'''

print("\n✅ Documentos gerados com sucesso!")
print(f"📄 {saida_informacoes}")
print(f"📄 {saida_portaria}")

# === Chamada do módulo web no final ===
try:
    acessar_web(dados)
    print(dados)
except Exception as e:
    print(f"\n⚠️ Falha ao preencher o sistema web: {e}")
